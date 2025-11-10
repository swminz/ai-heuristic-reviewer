import os
import io
import json
import zipfile
import base64
import datetime as dt

import streamlit as st
from PIL import Image, ImageDraw, ImageFont
from dotenv import load_dotenv
from openai import OpenAI

import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

load_dotenv()

# ===============================
# Helpers
# ===============================
def _to_base64(img: Image.Image) -> str:
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    return base64.b64encode(buf.getvalue()).decode("utf-8")

def chart_scores(heuristics_dict: dict) -> bytes:
    labels = ["H1","H2","H3","H4","H5","H6","H7","H8","H9","H10"]
    values = [heuristics_dict.get(h, 0) for h in labels]
    plt.figure()
    plt.bar(labels, values)           # (no custom colors)
    plt.ylim(0,3)
    plt.title("Heuristic Scores (0–3)")
    buf = io.BytesIO()
    plt.savefig(buf, format="png", bbox_inches="tight")
    plt.close()
    return buf.getvalue()

def chart_impact_effort(rows: list) -> bytes:
    m = {"Low":1,"Medium":2,"High":3}
    xs, ys, ids = [], [], []
    for r in rows or []:
        r = r or {}
        xs.append(m.get(r.get("effort","Medium"),2))
        ys.append(m.get(r.get("impact","Medium"),2))
        ids.append(r.get("issue_id",""))
    plt.figure()
    plt.scatter(xs, ys)
    for x,y,i in zip(xs,ys,ids): plt.text(x+0.03, y+0.03, i)
    plt.xticks([1,2,3], ["Low","Med","High"])
    plt.yticks([1,2,3], ["Low","Med","High"])
    plt.xlabel("Effort"); plt.ylabel("Impact")
    plt.title("Impact vs Effort")
    buf = io.BytesIO()
    plt.savefig(buf, format="png", bbox_inches="tight")
    plt.close()
    return buf.getvalue()

def _as_list(x):
    if x is None: return []
    if isinstance(x, list): return x
    if isinstance(x, dict): return [x]
    if isinstance(x, str): return [x]
    return []

def _parse_bbox(b):
    if isinstance(b, list) and len(b) >= 4: return b[:4]
    if isinstance(b, str):
        try:
            parts = [float(t) for t in b.strip().strip("[]()").split(",")]
            return parts[:4] if len(parts) >= 4 else None
        except Exception:
            return None
    return None

def _strip_code_fences(s: str) -> str:
    """Remove ```json ... ``` fences if the model adds them."""
    if not isinstance(s, str): return s
    s = s.strip()
    if s.startswith("```"):
        lines = s.splitlines()
        if lines and lines[0].startswith("```"): lines = lines[1:]
        if lines and lines[-1].strip().startswith("```"): lines = lines[:-1]
        s = "\n".join(lines).strip()
    return s

# -------- DOCX helpers --------
def _docx_add_bullets(doc, items):
    for it in (items or []):
        doc.add_paragraph(str(it), style='List Bullet')

def _docx_table_scores(doc, heuristics):
    table = doc.add_table(rows=1, cols=2)
    hdr = table.rows[0].cells
    hdr[0].text = "Heuristic"
    hdr[1].text = "Score (0–3)"
    labels = ["H1","H2","H3","H4","H5","H6","H7","H8","H9","H10"]
    names = [
        "Visibility of system status","Match between system & real world","User control & freedom",
        "Consistency & standards","Error prevention","Recognition vs recall",
        "Flexibility & efficiency","Aesthetic & minimalist","Error handling","Help & documentation"
    ]
    for lab, name in zip(labels, names):
        row = table.add_row().cells
        row[0].text = name
        row[1].text = str(heuristics.get(lab, ""))

def _build_docx(data, scores_png, ie_png, annotated_images_bytes, owner, date_str):
    doc = Document()
    title = (data.get("meta", {}) or {}).get("title") or "Heuristic Review"
    h = doc.add_heading(title, level=0); h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.add_paragraph(f"Owner: {owner} | Date: {date_str} | Version: v1")

    # Executive Summary
    doc.add_heading("Executive Summary", level=1)
    summary = (data.get("meta", {}) or {}).get("summary", {}) or {}
    doc.add_paragraph("Top issues:"); _docx_add_bullets(doc, summary.get("top_issues"))
    doc.add_paragraph("Quick wins:"); _docx_add_bullets(doc, summary.get("quick_wins"))
    doc.add_paragraph("Assumptions:"); _docx_add_bullets(doc, data.get("assumptions"))

    # Scorecard
    doc.add_heading("Heuristics Scorecard", level=1)
    heuristics = ((data.get("meta", {}) or {}).get("scores", {}) or {}).get("heuristics", {}) or {}
    _docx_table_scores(doc, heuristics)

    # Charts
    doc.add_heading("Charts", level=1)
    if scores_png:
        doc.add_paragraph("Heuristic scores")
        doc.add_picture(io.BytesIO(scores_png), width=Inches(6))
    if ie_png:
        doc.add_paragraph("Impact vs Effort")
        doc.add_picture(io.BytesIO(ie_png), width=Inches(6))

    # Annotated Screens
    doc.add_heading("Annotated Screens", level=1)
    for sid, png in annotated_images_bytes or []:
        doc.add_paragraph(sid)
        doc.add_picture(io.BytesIO(png), width=Inches(6))

    # Detailed Findings
    doc.add_heading("Detailed Findings", level=1)
    for iss in _as_list(data.get("issues")):
        if not isinstance(iss, dict): 
            continue
        doc.add_heading(f"{iss.get('id','ISS')} — {iss.get('title','(title)')} (Severity: {iss.get('severity','')})", level=2)
        doc.add_paragraph(f"Heuristic: {iss.get('heuristic','')}")
        doc.add_paragraph(f"Evidence: {iss.get('evidence','')}")
        rec = iss.get("recommendation", {}) or {}
        action = rec.get("action","") if isinstance(rec, dict) else (rec if isinstance(rec, str) else "")
        doc.add_paragraph(f"Recommendation: {action}")
        ac = _as_list(rec.get("acceptance_criteria") if isinstance(rec, dict) else [])
        if ac:
            doc.add_paragraph("Acceptance criteria:")
            for c in ac:
                doc.add_paragraph(str(c), style='List Bullet')

    # Impact/Effort
    doc.add_heading("Impact vs Effort (IDs)", level=1)
    for row in _as_list(data.get("impact_effort")):
        doc.add_paragraph(f"{row.get('issue_id')} [{row.get('impact')}/{row.get('effort')}]")

    # Appendix
    doc.add_heading("Appendix", level=1)
    screens = _as_list((data.get("meta", {}) or {}).get("screens"))
    doc.add_paragraph(f"Inputs: {len(screens)} screen(s)")
    doc.add_paragraph("Method & rubric: Nielsen + A11y spot checks")
    doc.add_paragraph("Limitations: Automated analysis; confirm with user testing where possible.")
    return doc

# ===============================
# CONFIG (OpenAI SDK v1+)
# ===============================
st.set_page_config(page_title="AI Heuristic Reviewer", page_icon="🕵️‍♀️", layout="wide")

OPENAI_API_KEY = st.secrets.get("OPENAI_API_KEY") or os.getenv("OPENAI_API_KEY")
DEFAULT_MODEL = st.secrets.get("OPENAI_VISION_MODEL") or os.getenv("OPENAI_VISION_MODEL","gpt-4o")
ORG = st.secrets.get("OPENAI_ORG")

if not OPENAI_API_KEY:
    st.error("❌ OPENAI_API_KEY is missing. Add it in Streamlit → Advanced settings → Secrets.")
    st.stop()

client_kwargs = {"api_key": OPENAI_API_KEY}
if ORG: client_kwargs["organization"] = ORG
client = OpenAI(**client_kwargs)

# ===============================
# Rubric text & JSON prompt
# ===============================
HEURISTIC_RUBRIC = r"""
Core Heuristics (0–3):
1. Visibility of system status
2. Match between system & real world
3. User control & freedom
4. Consistency & standards
5. Error prevention
6. Recognition vs recall (memory load)
7. Flexibility & efficiency
8. Aesthetic & minimalist design
9. Error handling (recognize/diagnose/recover)
10. Help & documentation

UX Quality (0–3): IA, Interaction clarity, Content, Cognitive load
A11y spot-checks (pass/warn/fail): Contrast (AA), focus order, targets ≥44×44, labels, keyboard reachability, visible error text
Severity: 0=None, 1=Low, 2=Moderate, 3=Major
"""

REPORT_TEMPLATE_MD = r"""
# Heuristic Review — <Product / Flow>
Owner: {owner} | Date: {date} | Version: v1

## Executive Summary
- Top issues:
{top_issues}
- Quick wins:
{quick_wins}
- Assumptions:
{assumptions}

## Heuristics Scorecard
| Heuristic | Score (0–3) |
|---|---|
| Visibility of system status | {h1} |
| Match between system & real world | {h2} |
| User control & freedom | {h3} |
| Consistency & standards | {h4} |
| Error prevention | {h5} |
| Recognition vs recall | {h6} |
| Flexibility & efficiency | {h7} |
| Aesthetic & minimalist | {h8} |
| Error handling | {h9} |
| Help & documentation | {h10} |

## Charts
{charts}

## Annotated Screens
{annotated_images}

## Detailed Findings
{findings}

## Impact vs Effort (IDs)
- {impact_effort}

## Appendix
- Method & rubric: Nielsen + A11y spot checks
- Inputs: {num_screens} screen(s)
- Limitations: Automated analysis; confirm with user testing where possible.
"""

SYSTEM_PROMPT_JSON = f"""
You are a Senior UX Evaluator. You will be given 1–5 UI screenshots (as image URLs) plus brief context.

Return **JSON only** matching this schema (keys & types), no extra text:

{{
  "meta": {{
    "title": "Heuristic Review — <screen/product>",
    "date": "<YYYY-MM-DD>",
    "screens": [{{"id":"screen_1","filename":"<name>","resolution":[W,H]}}],
    "summary": {{
      "top_issues": [str],
      "quick_wins": [str]
    }},
    "scores": {{
      "heuristics": {{"H1":int,"H2":int,"H3":int,"H4":int,"H5":int,"H6":int,"H7":int,"H8":int,"H9":int,"H10":int}},
      "ux_quality": {{"IA":int,"Interaction":int,"Content":int,"CognitiveLoad":int}},
      "a11y": {{"contrast":str,"focus":str,"targets":str,"labels":str}}
    }}
  }},
  "issues": [{{"id": "ISS-001","screen_id":"screen_1","title":str,"heuristic":"H<n> <name>","severity":int,"evidence":str,"bbox":[x,y,w,h],"recommendation":{{"action":str,"rationale":str,"acceptance_criteria":[str]}}}}],
  "impact_effort": [{{"issue_id":"ISS-001","impact":"Low|Medium|High","effort":"Low|Medium|High"}}],
  "assumptions": [str]
}}

CONSTRAINTS:
- Be specific; tie each issue to UI elements and include bbox.
- If little is wrong, still produce at least 3 concrete observations.
- Scores are 0–3. Severity is 1..3.
- Output must be valid JSON, no Markdown.
RUBRIC:
{HEURISTIC_RUBRIC}
"""

# ===============================
# UI
# ===============================
st.title("🕵️‍♀️ AI Heuristic Reviewer")
st.caption("Upload 1–5 screenshots. Get a professional heuristic review with annotated callouts, charts, and an impact/effort list.")

with st.sidebar:
    owner = st.text_input("Your name (for report)", value="Swati Minz")
    context = st.text_area("Context (optional)", placeholder="e.g., IRCTC search results; goal = choose a train & continue")
    model_choice = st.selectbox("Model", ["gpt-4o","gpt-4o-mini"], index=0 if DEFAULT_MODEL=="gpt-4o" else 1)
    submit_btn = st.button("Run Heuristic Review", type="primary")

uploads = st.file_uploader("Upload 1–5 screenshots (PNG/JPG)", type=["png","jpg","jpeg"], accept_multiple_files=True)

if uploads:
    st.subheader("Previews")
    for i, f in enumerate(uploads, start=1):
        img = Image.open(f).convert("RGB")
        st.image(img, caption=f"Screen {i}: {f.name}", use_column_width=True)

# ===============================
# Main
# ===============================
if submit_btn and uploads:
    today = dt.date.today().isoformat()
    local_images = []

    # Build JSON inputs (context + per-screen text + image_url)
    json_inputs = [{"type": "text", "text": f"Overall context: {context or '(none)'} | Date: {today}"}]
    for i, f in enumerate(uploads, start=1):
        f.seek(0); b = f.read()
        img = Image.open(io.BytesIO(b)).convert("RGB")
        w, h = img.size
        sid = f"screen_{i}"
        local_images.append((sid, f.name, img, w, h))
        json_inputs.append({"type": "text", "text": f"{sid}: {f.name}, resolution {w}x{h}"})
        json_inputs.append({"type": "image_url", "image_url": {"url": "data:image/png;base64," + _to_base64(img)}})

    # --------- JSON call ----------
    with st.spinner("Analyzing with OpenAI (JSON)…"):
        try:
            resp = client.chat.completions.create(
                model=model_choice,
                messages=[
                    {"role": "system", "content": SYSTEM_PROMPT_JSON},
                    {"role": "user", "content": json_inputs}
                ],
                temperature=0.1,
                max_tokens=4000,
                response_format={"type": "json_object"},  # force strict JSON
            )
            raw_json = resp.choices[0].message.content or ""
            raw_json = _strip_code_fences(raw_json)
        except Exception as e:
            msg = str(e)
            if "insufficient_quota" in msg or "Error code: 429" in msg:
                st.error("Your API key has **no available quota**. Add billing/credits at platform.openai.com → Billing.")
            else:
                st.error(f"OpenAI error: {e}")
            st.stop()

    try:
        data = json.loads(raw_json)
    except Exception:
        st.error("The model did not return valid JSON. Raw response shown below.")
        st.code(raw_json)
        st.stop()

    meta = data.get("meta", {}) or {}
    summary = meta.get("summary", {}) or {}
    scores = meta.get("scores", {}) or {}
    H = scores.get("heuristics", {}) if isinstance(scores, dict) else {}
    issues = _as_list(data.get("issues"))
    impact_effort = _as_list(data.get("impact_effort"))

    if len(issues) < 3:
        st.warning("The model returned very few findings. Try uploading one more full-page screen.")

    # -------- Annotate images & ZIP --------
    annotated_images_md = []
    annotated_images_bytes = []  # keep for DOCX
    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zf:
        screen_map = {sid: (fname, im.copy(), w, h) for (sid, fname, im, w, h) in local_images}
        grouped = {}
        for iss in issues:
            if isinstance(iss, dict):
                grouped.setdefault(iss.get("screen_id","screen_1"), []).append(iss)

        font = None
        try:
            font = ImageFont.truetype("arial.ttf", 20)
        except Exception:
            pass

        for sid, items in grouped.items():
            if sid not in screen_map and screen_map:
                sid = list(screen_map.keys())[0]
            fname, base_img, w, h = screen_map[sid]
            draw = ImageDraw.Draw(base_img)

            for idx, iss in enumerate(items, start=1):
                bbox = _parse_bbox(iss.get("bbox"))
                if bbox:
                    x,y,bw,bh = bbox
                    x0,y0 = int(x*w), int(y*h)
                    x1,y1 = int((x+bw)*w), int((y+bh)*h)
                    draw.rectangle([x0,y0,x1,y1], outline=(255,0,0), width=3)
                    tx,ty = x0, max(0, y0-24)
                    draw.rectangle([tx,ty,tx+24,ty+24], fill=(255,0,0))
                    draw.text((tx+7,ty+4), f"{idx}", fill=(255,255,255), font=font)

            buf = io.BytesIO()
            base_img.save(buf, format="PNG")
            png_bytes = buf.getvalue()
            zf.writestr(f"images/{sid}_annotated.png", png_bytes)
            annotated_images_md.append(f"![{sid} annotated](images/{sid}_annotated.png)")
            annotated_images_bytes.append((sid, png_bytes))

        # Originals
        for i, (sid, fname, im, _, _) in enumerate(local_images, start=1):
            buf = io.BytesIO(); im.save(buf, format="PNG")
            zf.writestr(f"images/original_{i}_{fname}", buf.getvalue())

    # -------- Charts --------
    try:
        scores_png = chart_scores(H or {})
        ie_png = chart_impact_effort(impact_effort)
        with zipfile.ZipFile(zip_buffer, 'a', zipfile.ZIP_DEFLATED) as zf:
            zf.writestr("charts/heuristics.png", scores_png)
            zf.writestr("charts/impact_effort.png", ie_png)
        charts_md = "![Heuristic scores](charts/heuristics.png)\n\n![Impact vs Effort](charts/impact_effort.png)"
    except Exception:
        scores_png = None
        ie_png = None
        charts_md = "(charts unavailable)"

    # -------- Summary & Findings (Markdown preview) --------
    def _bullets(x):
        items = _as_list(x)
        return "\n".join([f"  - {i}" for i in items]) or "  - (none)"

    top_issues_md = _bullets(summary.get("top_issues"))
    quick_wins_md = _bullets(summary.get("quick_wins"))
    assumptions_md = _bullets(data.get("assumptions"))

    def fmt_issue(iss: dict) -> str:
        if not isinstance(iss, dict):
            return "### (Unparsed issue)\n- Heuristic:\n- Evidence:\n- Recommendation:\n- Acceptance criteria:\n  - (none)\n"
        rec = iss.get("recommendation", {}) or {}
        action = ""
        ac_list = []
        if isinstance(rec, dict):
            action = rec.get("action","") or ""
            ac_list = _as_list(rec.get("acceptance_criteria"))
        elif isinstance(rec, list):
            ac_list = rec
        elif isinstance(rec, str):
            action = rec
        ac = "\n".join([f"  - {c}" for c in _as_list(ac_list)]) or "  - (none)"
        sev = iss.get("severity", 0)
        try: sev = int(sev)
        except Exception: sev = 0
        return f"""### {iss.get('id','ISS')} — {iss.get('title','(title)')} (Severity: {sev})
- Heuristic: {iss.get('heuristic','')}
- Evidence: {iss.get('evidence','')}
- Recommendation: {action}
- Acceptance criteria:
{ac}
"""

    findings_md = "\n".join([fmt_issue(i) for i in issues]) or "(No issues parsed)"
    impact_effort_md = ", ".join([f"{(ie or {}).get('issue_id')} [{(ie or {}).get('impact')}/{(ie or {}).get('effort')}]" for ie in impact_effort]) or "(n/a)"

    final_report = REPORT_TEMPLATE_MD.format(
        owner=owner, date=today,
        top_issues=top_issues_md, quick_wins=quick_wins_md, assumptions=assumptions_md,
        h1=(H.get("H1","")), h2=(H.get("H2","")), h3=(H.get("H3","")), h4=(H.get("H4","")),
        h5=(H.get("H5","")), h6=(H.get("H6","")), h7=(H.get("H7","")), h8=(H.get("H8","")),
        h9=(H.get("H9","")), h10=(H.get("H10","")),
        charts=charts_md,
        annotated_images="\n\n".join(annotated_images_md) or "(annotated images included in ZIP)",
        findings=findings_md, impact_effort=impact_effort_md,
        num_screens=len(local_images)
    )

    st.subheader("Report Preview")
    st.markdown(final_report)

    # ZIP (Markdown + charts + images + raw JSON)
    with zipfile.ZipFile(zip_buffer, 'a', zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("report/report_compiled.md", final_report)
        zf.writestr("report/data.json", json.dumps(data, indent=2))

    st.download_button("⬇️ Download assets (ZIP)", data=zip_buffer.getvalue(), file_name="heuristic_review.zip")

    # DOCX export (single file, easy to convert to PDF)
    try:
        doc = _build_docx(data, scores_png, ie_png, annotated_images_bytes, owner, today)
        docx_buf = io.BytesIO()
        doc.save(docx_buf); docx_buf.seek(0)
        st.download_button(
            "📄 Download Word report (DOCX)",
            data=docx_buf.getvalue(),
            file_name="heuristic_review.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    except Exception as e:
        st.warning(f"Could not build DOCX: {e}")

else:
    st.info("Upload at least one full-page screenshot, then click ‘Run Heuristic Review’. For best results, choose **Model = gpt-4o**.")
